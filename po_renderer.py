"""
PO renderer — turns a PO draft dict into a polished .docx file.

Base document is templates/po-letterhead.docx — a stripped copy of the
client's branded letterhead, providing the page header image, footer
contact text, page size, margins, and theme/styles. Body content is
appended programmatically by the _render_* functions.

If the template file is missing, falls back to a blank Document and
logs a warning — generation still works in dev/test environments but
the output will be unbranded.
"""
from __future__ import annotations

import io
import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)

_TEMPLATE_PATH = Path(__file__).parent / "templates" / "po-letterhead.docx"


# Brand palette (kept aligned with DESIGN.md corporate theme so the PO looks
# coherent with the dashboard's "Corporate" skin)
BRAND_NAVY = "1E3A5F"
BRAND_TEXT = "1A202C"
BRAND_MUTED = "64748B"
BRAND_BORDER = "E2E8F0"
BRAND_ACCENT = "D97706"  # subtle orange — sparingly


def _ensure_python_docx():
    """Import python-docx lazily so the rest of the app starts even if the
    dep is missing in dev environments."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return {
            "Document": Document, "Pt": Pt, "RGBColor": RGBColor,
            "Inches": Inches, "Cm": Cm,
            "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
            "WD_ALIGN_VERTICAL": WD_ALIGN_VERTICAL,
            "qn": qn, "OxmlElement": OxmlElement,
        }
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for PO rendering. "
            "Install with `pip install python-docx`."
        ) from exc


def _set_cell_shading(cell, hex_color: str):
    """Apply a background fill to a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_borders(cell, color: str = BRAND_BORDER, size: str = "4"):
    """Apply a uniform thin border to a cell. size is in 1/8 pt — 4 ≈ 0.5pt."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _format_money(amount: float, currency: str = "USD") -> str:
    """Format an amount as currency with thousands separators."""
    try:
        return f"{currency} {amount:,.2f}"
    except (TypeError, ValueError):
        return f"{currency} 0.00"


def _add_styled_run(paragraph, text: str, *, bold: bool = False, size: int = 10,
                     color: str = BRAND_TEXT, font: str = "Calibri"):
    """Append a run to a paragraph with the given styling. Returns the run."""
    from docx.shared import Pt, RGBColor
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def _add_horizontal_rule(doc, color: str = BRAND_NAVY):
    """Insert a horizontal rule using a single-row, single-cell table with a top border."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    cell.text = ""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:color"), color)
    tcBorders.append(top)
    for edge in ("left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _strip_leading_empty_paragraphs(doc) -> None:
    """Remove blank paragraphs at the top of the body so rendered content
    starts at the page top margin. Skips anything with runs/text/images —
    only truly empty paragraphs are dropped. Stops at the first non-empty
    element so any tables / styled content the template intends to keep
    are preserved."""
    while doc.paragraphs:
        p = doc.paragraphs[0]
        if p.text.strip() or p.runs:
            break
        p._element.getparent().remove(p._element)


def _render_title_block(doc, draft: dict[str, Any]):
    """The PURCHASE ORDER title + PO number + date in a two-cell layout."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells

    # Left: title
    left_p = left.paragraphs[0]
    _add_styled_run(left_p, "PURCHASE ORDER", bold=True, size=22, color=BRAND_NAVY)

    # Right: PO number + date stacked
    right_p1 = right.paragraphs[0]
    right_p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(right_p1, "PO Number  ", size=9, color=BRAND_MUTED)
    _add_styled_run(right_p1, draft.get("po_number", ""),
                    bold=True, size=12, color=BRAND_TEXT)

    right_p2 = right.add_paragraph()
    right_p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(right_p2, "Date  ", size=9, color=BRAND_MUTED)
    po_date = draft.get("po_date") or datetime.now(timezone.utc).strftime("%Y-%m-%d")
    _add_styled_run(right_p2, po_date, bold=True, size=11, color=BRAND_TEXT)

    # Remove all borders from this title-block table
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in (left, right):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            tcBorders.append(b)
        tcPr.append(tcBorders)


def _render_party_block(doc, draft: dict[str, Any]):
    """SUPPLIER block, plus SHIP TO when the incoterm makes a buyer-side
    delivery address meaningful. Under seller-side terms (FCA/FOB/EXW/FAS)
    the reconciler clears ship_to and we drop the column entirely — a
    PO that promises FCA at the supplier's dock and also lists a buyer
    SHIP TO is internally contradictory."""
    ship_to = draft.get("ship_to") or {}
    has_ship_to = bool(ship_to.get("name") or (ship_to.get("lines") or []))

    if has_ship_to:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        sup_cell, ship_cell = table.rows[0].cells
        rows = [
            (sup_cell, "SUPPLIER", draft.get("supplier") or {}),
            (ship_cell, "SHIP TO", ship_to),
        ]
    else:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        sup_cell = table.rows[0].cells[0]
        rows = [(sup_cell, "SUPPLIER", draft.get("supplier") or {})]

    for cell, title, party in rows:
        _set_cell_borders(cell)
        # Reset the auto-empty paragraph
        cell.text = ""
        p_title = cell.paragraphs[0]
        _add_styled_run(p_title, title, bold=True, size=9, color=BRAND_NAVY)

        if party.get("name"):
            p_name = cell.add_paragraph()
            _add_styled_run(p_name, party["name"], bold=True, size=11, color=BRAND_TEXT)

        for line in (party.get("lines") or []):
            if not line:
                continue
            p_line = cell.add_paragraph()
            _add_styled_run(p_line, line, size=10, color=BRAND_TEXT)

        # Supplier extras: attention + email + cage
        if title == "SUPPLIER":
            if party.get("attention"):
                p = cell.add_paragraph()
                _add_styled_run(p, "Attention: ", size=9, color=BRAND_MUTED)
                _add_styled_run(p, party["attention"], size=10, color=BRAND_TEXT)
            if party.get("email"):
                p = cell.add_paragraph()
                _add_styled_run(p, party["email"], size=9, color=BRAND_MUTED)
            if party.get("cage_code"):
                p = cell.add_paragraph()
                _add_styled_run(p, f"CAGE: {party['cage_code']}", size=9, color=BRAND_MUTED)


def _render_references(doc, draft: dict[str, Any]):
    """Compact references panel."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_borders(cell)
    _set_cell_shading(cell, "F8F9FA")
    cell.text = ""

    p_title = cell.paragraphs[0]
    _add_styled_run(p_title, "REFERENCES", bold=True, size=9, color=BRAND_NAVY)

    def _ref_line(label: str, value: str):
        if not value:
            return
        p = cell.add_paragraph()
        _add_styled_run(p, f"{label}  ", size=9, color=BRAND_MUTED)
        _add_styled_run(p, value, size=10, color=BRAND_TEXT)

    _ref_line("Contract", draft.get("contract_no", ""))
    _ref_line("Contract Date", draft.get("contract_date", ""))
    _ref_line("Client Reference", draft.get("client_reference", ""))
    _ref_line("Supplier Quote", draft.get("quote_no", ""))
    _ref_line("Quote Date", draft.get("quote_date", ""))
    title = draft.get("title", "")
    if title:
        _ref_line("Description", title)


def _render_items_table(doc, draft: dict[str, Any]):
    """Line items table."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    p = doc.add_paragraph()
    _add_styled_run(p, "PURCHASE DETAILS", bold=True, size=10, color=BRAND_NAVY)

    currency = draft.get("currency", "USD")
    headers = ["Line", "Part Number", "Description", "Qty", "Unit", f"Unit Price ({currency})", f"Extended ({currency})"]
    items = draft.get("items") or []

    table = doc.add_table(rows=1 + len(items) + 2, cols=len(headers))
    table.autofit = True

    # Header row
    header_row = table.rows[0]
    for col_idx, label in enumerate(headers):
        cell = header_row.cells[col_idx]
        cell.text = ""
        _set_cell_borders(cell, color=BRAND_NAVY, size="6")
        _set_cell_shading(cell, BRAND_NAVY)
        p = cell.paragraphs[0]
        if col_idx in (3, 5, 6):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_styled_run(p, label, bold=True, size=9, color="FFFFFF")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Item rows
    for row_idx, item in enumerate(items, start=1):
        row = table.rows[row_idx]
        desc = item.get("description", "")
        nsn = item.get("nsn", "")
        if nsn:
            desc = f"{desc}  (NSN: {nsn})" if desc else f"NSN: {nsn}"
        # Defensive numeric coercion — never trust types coming from JSON.
        # validate_draft_math already caught bad values, but if the renderer
        # is ever called outside that path we still want a clean string.
        try:
            qty_val = float(item.get("qty", 0) or 0)
        except (TypeError, ValueError):
            qty_val = 0.0
        qty_str = str(int(qty_val)) if qty_val.is_integer() else f"{qty_val:g}"
        try:
            unit_price_val = float(item.get("unit_price", 0) or 0)
        except (TypeError, ValueError):
            unit_price_val = 0.0
        try:
            extended_val = float(item.get("extended", 0) or 0)
        except (TypeError, ValueError):
            extended_val = 0.0
        values = [
            str(item.get("line", row_idx)),
            item.get("part_no", "") or "",
            desc,
            qty_str,
            item.get("unit", "EA"),
            f"{unit_price_val:,.2f}",
            f"{extended_val:,.2f}",
        ]
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = ""
            _set_cell_borders(cell)
            p = cell.paragraphs[0]
            if col_idx in (3, 5, 6):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_styled_run(p, str(value), size=10, color=BRAND_TEXT)

        # Optional second-line lead-time note
        if item.get("lead_time"):
            note_row_idx = row_idx
            # Append lead time as a small italic line *inside* the description cell.
            cell = row.cells[2]
            lt_p = cell.add_paragraph()
            _add_styled_run(lt_p, f"Lead time: {item['lead_time']}",
                            size=8, color=BRAND_MUTED)

    # Subtotal row (second-to-last)
    subtotal_row = table.rows[1 + len(items)]
    for i in range(len(headers)):
        cell = subtotal_row.cells[i]
        cell.text = ""
        _set_cell_borders(cell)
    # Merge cells 0..5 for the label
    label_cell = subtotal_row.cells[0]
    for j in range(1, 6):
        label_cell = label_cell.merge(subtotal_row.cells[j])
    label_cell.text = ""
    p_lbl = label_cell.paragraphs[0]
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(p_lbl, "Subtotal", bold=True, size=10, color=BRAND_TEXT)
    val_cell = subtotal_row.cells[6]
    val_cell.text = ""
    p_val = val_cell.paragraphs[0]
    p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(p_val, f"{float(draft.get('subtotal', 0)):,.2f}",
                    bold=True, size=10, color=BRAND_TEXT)

    # Total row (last)
    total_row = table.rows[2 + len(items)]
    for i in range(len(headers)):
        cell = total_row.cells[i]
        cell.text = ""
        _set_cell_borders(cell, color=BRAND_NAVY, size="6")
        _set_cell_shading(cell, "F1F5F9")
    label_cell = total_row.cells[0]
    for j in range(1, 6):
        label_cell = label_cell.merge(total_row.cells[j])
    label_cell.text = ""
    p_lbl = label_cell.paragraphs[0]
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(p_lbl, f"TOTAL ({currency})", bold=True, size=11, color=BRAND_NAVY)
    val_cell = total_row.cells[6]
    val_cell.text = ""
    p_val = val_cell.paragraphs[0]
    p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(p_val, f"{float(draft.get('total', 0)):,.2f}",
                    bold=True, size=12, color=BRAND_NAVY)

    # Tax note below
    tax_p = doc.add_paragraph()
    _add_styled_run(tax_p, draft.get("tax_note") or "Taxes: As Applicable",
                    size=9, color=BRAND_MUTED)


def _render_terms_block(doc, draft: dict[str, Any]):
    """Delivery / incoterms / payment-terms row. Cells with no value are
    omitted entirely — emitting a header like `INCOTERMS` with a `—`
    placeholder is misleading on a binding PO."""
    fields = [
        ("DELIVERY BY", (draft.get("delivery_date") or "").strip()),
        ("INCOTERMS", (draft.get("incoterms") or "").strip()),
        ("PAYMENT TERMS", (draft.get("payment_terms") or "").strip()),
    ]
    populated = [(label, value) for label, value in fields if value]
    if not populated:
        return

    table = doc.add_table(rows=1, cols=len(populated))
    table.autofit = True
    for cell, (label, value) in zip(table.rows[0].cells, populated):
        _set_cell_borders(cell)
        _set_cell_shading(cell, "F8F9FA")
        cell.text = ""
        p_lbl = cell.paragraphs[0]
        _add_styled_run(p_lbl, label, bold=True, size=8, color=BRAND_MUTED)
        p_val = cell.add_paragraph()
        _add_styled_run(p_val, value, size=10, color=BRAND_TEXT)


def _render_flow_down(doc, draft: dict[str, Any]):
    """Contract flow-down clauses bullet list. Always starts on a new page
    so the supplier terms read as a distinct contractual section, separate
    from the price/quantity/incoterm header on page 1. The packaging spec
    expansion that follows can continue on the same page — just a spacer
    between them (handled in _render_packaging_expansion).

    Uses a literal "•" prefix rather than the "List Bullet" paragraph
    style so rendering doesn't depend on Word's built-in list styles
    being present (the stripped letterhead template drops unused styles)."""
    from docx.shared import Pt

    clauses = draft.get("flow_down_clauses") or []
    if not clauses:
        return

    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True  # heading stays with first bullet
    _add_styled_run(p, "FLOW-DOWN CONTRACT TERMS",
                    bold=True, size=12, color=BRAND_NAVY)
    for clause in clauses:
        bullet = doc.add_paragraph()
        bullet.paragraph_format.left_indent = bullet.paragraph_format.left_indent or None
        _add_styled_run(bullet, "•  ", size=9, color=BRAND_MUTED)
        _add_styled_run(bullet, clause, size=9, color=BRAND_TEXT)


# DND packaging spec expansions. When the flow-down clauses reference any of
# these specs, render the expanded plain-English requirements as a separate
# block so the supplier doesn't have to look up the underlying standards.
# Matches the client reference PO's structure exactly.
_PACKAGING_EXPANSIONS = {
    "D-LM-008-036/SF-000": [
        "All item numbers shall be prepared for delivery in accordance with the latest issue of Canadian Forces Packaging Specification D‑LM‑008‑036/SF‑000.",
        "Each item shall be packaged in quantities of one (1) item per package.",
        "Required markings shall comply with D2000C (Markings).",
        "Required labels shall comply with D2001C (Labelling).",
        "Wood packaging materials shall comply with D2025C requirements.",
        "Packaging shall protect against handling, storage and transport damage and preserve product integrity throughout shipment.",
    ],
    "D2000C": [
        "D2000C – Marking Requirements",
        "Exterior shipping containers shall be clearly marked with part number, nomenclature, quantity, contract/PO number, and shipment identification.",
        "Markings shall be durable, legible, and placed in visible locations on the package exterior.",
        "Required handling and storage markings shall be applied where applicable.",
    ],
    "D2001C": [
        "D2001C – Labelling Requirements",
        "Labels shall be securely affixed and readable throughout handling and transportation.",
        "Labels shall identify contents and any required logistics information.",
        "Labels shall remain intact and resistant to normal shipping and storage conditions.",
    ],
    "D2025C": [
        "D2025C – Wood Packaging Material Requirements",
        "Any wood packaging material used shall comply with applicable international phytosanitary requirements.",
        "Wood packaging shall be heat-treated or otherwise compliant where required.",
        "Required certification or markings for treated wood packaging shall be visible and maintained.",
    ],
}


def _render_packaging_expansion(doc, draft: dict[str, Any]):
    """If the flow-down clauses reference DND packaging specs, append an
    expanded plain-English explanation block — the supplier shouldn't need
    to look up the underlying standards. Matches the client reference PO.

    Each spec is expanded only once even if referenced multiple times.
    Order is: top-level packaging spec first, then the sub-specs.

    Packaging continues on the same page as the flow-down terms (no
    forced page break — that's on FLOW-DOWN itself). A visible spacer
    before the title separates it from the flow-down bullets. Word can
    flow the subsections naturally; keep_with_next on each subsection
    heading prevents the heading from getting orphaned at the bottom
    of a page if a natural break does land mid-block."""
    from docx.shared import Pt

    clauses_blob = " ".join(draft.get("flow_down_clauses") or [])
    if not clauses_blob:
        return

    matched: list[tuple[str, list[str]]] = []
    for key in ("D-LM-008-036/SF-000", "D2000C", "D2001C", "D2025C"):
        if key in clauses_blob:
            matched.append((key, _PACKAGING_EXPANSIONS[key]))
    if not matched:
        return

    title = doc.add_paragraph()
    # Vertical breathing room between the flow-down bullets above and
    # this section header — no page break, just a clear visual gap.
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    _add_styled_run(title, "Detailed Packaging Requirements",
                    bold=True, size=12, color=BRAND_NAVY)

    for key, lines in matched:
        # First line is the section heading (e.g. "D2000C – Marking
        # Requirements"); render bold. Remaining lines are bullets.
        if not lines:
            continue
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(4)
        head.paragraph_format.keep_with_next = True  # don't orphan heading
        _add_styled_run(head, lines[0], bold=True, size=9, color=BRAND_TEXT)
        for i, line in enumerate(lines[1:]):
            bullet = doc.add_paragraph()
            # Tie every bullet except the last one to the next paragraph,
            # so a subsection won't split mid-list. The last bullet is
            # released so Word can break cleanly between subsections.
            is_last_in_subsection = (i == len(lines) - 2)
            if not is_last_in_subsection:
                bullet.paragraph_format.keep_with_next = True
            _add_styled_run(bullet, "•  ", size=9, color=BRAND_MUTED)
            _add_styled_run(bullet, line, size=9, color=BRAND_TEXT)


def _render_notes(doc, draft: dict[str, Any]):
    notes = draft.get("notes", "")
    if not notes:
        return
    p = doc.add_paragraph()
    _add_styled_run(p, "NOTES", bold=True, size=10, color=BRAND_NAVY)
    body = doc.add_paragraph()
    _add_styled_run(body, notes, size=10, color=BRAND_TEXT)


def _render_signature(doc, draft: dict[str, Any]):
    """Authorized-by block with a physical signature line above the printed
    name, matching the client reference PO."""
    doc.add_paragraph()  # spacer
    _add_horizontal_rule(doc, color=BRAND_BORDER)

    auth = draft.get("authorized_by") or {}
    p1 = doc.add_paragraph()
    _add_styled_run(p1, "Authorized by:  ", size=10, color=BRAND_TEXT)
    # Underscore line for ink signature. Width tuned to roughly match the
    # client reference PO (~35 chars of "_" reads as a ~3.5" line at 10pt).
    _add_styled_run(p1, "_" * 35, size=10, color=BRAND_TEXT)
    p2 = doc.add_paragraph()
    _add_styled_run(p2, auth.get("name", ""), bold=True, size=11, color=BRAND_TEXT)
    p3 = doc.add_paragraph()
    _add_styled_run(p3, auth.get("title", ""), size=10, color=BRAND_MUTED)


# ── Public API ────────────────────────────────────────────────────────────────

def render(draft: dict[str, Any]) -> bytes:
    """Render a PO draft to a .docx file. Returns the file bytes.

    Raises RuntimeError if python-docx is not installed.
    """
    mods = _ensure_python_docx()
    Document = mods["Document"]

    # Load the branded letterhead as the base — header image, footer
    # contact block, page size, margins, and theme all come along.
    if _TEMPLATE_PATH.is_file():
        doc = Document(str(_TEMPLATE_PATH))
        # Template ships with a few empty paragraphs as breathing room.
        # Drop them so our title block sits at the top margin; the page
        # header (letterhead image) is rendered by Word regardless and
        # is unaffected by body paragraphs.
        _strip_leading_empty_paragraphs(doc)
    else:
        log.warning(
            "PO letterhead template not found at %s — falling back to blank document",
            _TEMPLATE_PATH,
        )
        doc = Document()

    _render_title_block(doc, draft)
    doc.add_paragraph()
    _render_party_block(doc, draft)
    doc.add_paragraph()
    _render_references(doc, draft)
    doc.add_paragraph()
    _render_items_table(doc, draft)
    doc.add_paragraph()
    _render_terms_block(doc, draft)
    # No spacer paragraph here — _render_flow_down's title uses
    # page_break_before to start on a fresh page. A trailing empty
    # paragraph at the end of page 1 was getting pushed onto page 2
    # by the page break, producing a near-blank middle page.
    _render_flow_down(doc, draft)
    _render_packaging_expansion(doc, draft)
    _render_notes(doc, draft)
    _render_signature(doc, draft)

    buf = io.BytesIO()
    doc.save(buf)
    log.info(
        "po.render po_number=%r items=%d bytes=%d",
        draft.get("po_number"), len(draft.get("items") or []), buf.tell(),
    )
    return buf.getvalue()
